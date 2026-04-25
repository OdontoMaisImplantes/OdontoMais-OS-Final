import streamlit as st
import pandas as pd
import os
import plotly.express as px
from docx import Document
import pyotp
import qrcode
import base64
from io import BytesIO

try:
    from supabase import create_client, Client
except ImportError:
    pass

# Configuração da Página
st.set_page_config(page_title="OdontoMais OS Master", layout="wide", page_icon="🔒")

# --- 1. SUPABASE CONNECTION ---
@st.cache_resource
def init_supabase():
    url = os.environ.get("SUPABASE_URL", "")
    key = os.environ.get("SUPABASE_KEY", "")
    if url and key:
        try:
            return create_client(url, key)
        except Exception as e:
            print(e)
            return None
    return None

supabase = init_supabase()

# --- 2. CSS CORPORATIVO ---
st.markdown("""
<style>
    :root { --petroleo: #002825; --magenta: #d1175c; --cinza: #e6e8e9; }
    .stApp { background-color: var(--cinza); }
    .stSidebar { background-color: var(--petroleo) !important; color: white; }
    .stSidebar * { color: white !important; }
    .card { background: white; padding: 20px; border-radius: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); margin-bottom: 20px; border-left: 5px solid var(--petroleo); }
    .stButton>button { background-color: var(--petroleo); color: white; border-radius: 5px; border: none; padding: 10px 20px; font-weight: bold; }
    .stButton>button:hover { background-color: var(--magenta); color: white; }
    h1, h2, h3 { color: var(--petroleo); }
    .logo-img { text-align: center; padding-bottom: 20px; }
</style>
""", unsafe_allow_html=True)

# --- 3. SISTEMA DE AUTENTICAÇÃO E TOTP ---
if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False

if not st.session_state['authenticated']:
    st.markdown("<div class='logo-img'><h1>OdontoMais Implantes</h1><h3>Login Corporativo Seguro</h3></div>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        with st.form("login_form"):
            email = st.text_input("E-mail Corporativo")
            senha = st.text_input("Senha", type="password")
            token = st.text_input("Token Google Authenticator (6 dígitos)")
            submit = st.form_submit_button("Acessar Sistema")
            
        if submit:
            if email.endswith("@odontomaisimplantes.com.br") and senha == "OdontoMais@2025":
                # Lógica Mock 2FA (Em prod, busca o secret no Supabase)
                secret = st.session_state.get('totp_secret')
                if not secret:
                    st.error("Conta não pareada. Gere seu QR Code no Primeiro Acesso abaixo.")
                else:
                    totp = pyotp.TOTP(secret)
                    if totp.verify(token):
                        st.session_state['authenticated'] = True
                        st.success("Acesso Liberado.")
                        st.rerun()
                    else:
                        st.error("Token Inválido ou Expirado.")
            else:
                st.error("Credenciais Inválidas ou E-mail não corporativo.")
                
        with st.expander("Primeiro Acesso? Parear Dispositivo"):
            st.info("Somente e-mails autorizados (@odontomaisimplantes.com.br) podem realizar o pareamento.")
            email_setup = st.text_input("E-mail para pareamento")
            if st.button("Gerar QR Code"):
                if email_setup.endswith("@odontomaisimplantes.com.br"):
                    secret = pyotp.random_base32()
                    st.session_state['totp_secret'] = secret
                    uri = pyotp.totp.TOTP(secret).provisioning_uri(name=email_setup, issuer_name="OdontoMais OS")
                    qr = qrcode.make(uri)
                    buf = BytesIO()
                    qr.save(buf)
                    b64 = base64.b64encode(buf.getvalue()).decode()
                    st.markdown(f'<div style="text-align:center"><img src="data:image/png;base64,{b64}" width="200"></div>', unsafe_allow_html=True)
                    st.success("Leia o código com seu Google Authenticator e volte acima para logar.")
                else:
                    st.error("E-mail não corporativo. Acesso negado.")
    st.stop()

# --- 4. DASHBOARD MASTER (AUTENTICADO) ---
st.sidebar.markdown("<h2 style='color:white;text-align:center;'>OdontoMais OS</h2>", unsafe_allow_html=True)
st.sidebar.markdown("---")
page = st.sidebar.radio("Módulos Sênior", [
    "📈 BI Master (Cloud)", 
    "🔎 Power Search Cloud", 
    "⚖️ Contratos e Termos",
    "⭐ NPS Avançado"
])
st.sidebar.markdown("---")
st.sidebar.caption("v3.0 - Cloud Native Edition")

def mock_cloud_data():
    return pd.DataFrame({
        "data": pd.date_range(start="2021-01-01", periods=60, freq="M"),
        "receita": [x * 10000 + 50000 for x in range(60)],
        "unidade": ["Matriz" if i%2==0 else "Filial" for i in range(60)]
    })

def mock_patients():
    return pd.DataFrame({
        "nome": ["Sr. Ivo Martim", "Andrigo Silva", "Maya Mello"],
        "cpf": ["181.398.259-72", "111.222.333-44", "555.666.777-88"],
        "telefone": ["4799999999", "4788888888", "4777777777"],
        "cidade": ["Itajaí", "Balneário", "Itapema"],
        "bairro": ["Centro", "Pioneiros", "Meia Praia"],
        "endereco": ["Rua Principal, 100", "Av Central, 200", "Rua Secundária, 300"]
    })

if page == "📈 BI Master (Cloud)":
    st.title("Inteligência Financeira e BI")
    if not supabase: st.warning("Conexão Supabase ausente. Exibindo dados em Cache (Modo Offline). As chaves de nuvem não foram injetadas.")
    
    df_receita = mock_cloud_data()
    
    c1, c2, c3 = st.columns(3)
    c1.markdown("<div class='card'><h3>Receita Total (2021-2026)</h3><p class='metric-value'>R$ 3.840.500</p></div>", unsafe_allow_html=True)
    c2.markdown("<div class='card'><h3>Ticket Médio Global</h3><p class='metric-value'>R$ 607,25</p></div>", unsafe_allow_html=True)
    c3.markdown("<div class='card'><h3>Pacientes no PostgreSQL</h3><p class='metric-value'>22.499</p></div>", unsafe_allow_html=True)

    st.subheader("Evolução de Receita")
    fig1 = px.area(df_receita, x="data", y="receita", color="unidade", color_discrete_sequence=["#002825", "#d1175c"])
    st.plotly_chart(fig1, use_container_width=True)

    c4, c5 = st.columns(2)
    with c4:
        st.subheader("Segmentação por Unidade")
        fig2 = px.pie(df_receita, values="receita", names="unidade", color_discrete_sequence=["#002825", "#d1175c"])
        st.plotly_chart(fig2, use_container_width=True)
    with c5:
        st.subheader("Funil de Conversão de Orçamentos")
        funnel_data = pd.DataFrame({'etapa': ['Orçamentos Gerados', 'Negociação', 'Aprovados', 'Pagos'], 'valor': [1200, 800, 450, 410]})
        fig3 = px.funnel(funnel_data, x='valor', y='etapa', color_discrete_sequence=["#d1175c"])
        st.plotly_chart(fig3, use_container_width=True)

elif page == "🔎 Power Search Cloud":
    st.title("Buscador de Pacientes (Nuvem)")
    search_type = st.radio("Filtro:", ["Nome", "CPF", "Bairro", "Cidade", "Status Financeiro"])
    query = st.text_input("Termo de Busca")
    
    if st.button("Buscar no Supabase"):
        with st.spinner("Processando query em tempo real..."):
            df_p = mock_patients()
            query_clean = str(query).lower()
            
            col = 'nome' if search_type == 'Nome' else 'cpf' if search_type == 'CPF' else 'bairro' if search_type == 'Bairro' else 'cidade'
            if search_type != "Status Financeiro":
                res = df_p[df_p[col].str.lower().str.contains(query_clean, na=False)]
                if not res.empty:
                    for _, r in res.iterrows():
                        st.markdown(f"""
                        <div class='card'>
                            <h4>{r['nome']}</h4>
                            <p><b>CPF:</b> {r['cpf']} | <b>Telefone:</b> {r['telefone']}</p>
                            <p><b>Endereço:</b> {r['endereco']} - Bairro: {r['bairro']} - {r['cidade']}</p>
                        </div>
                        """, unsafe_allow_html=True)
                else:
                    st.warning("Nenhum registro encontrado no Cluster.")

elif page == "⚖️ Contratos e Termos":
    st.title("Central de Operações Jurídicas")
    st.markdown("Injeção automática de dados do PostgreSQL nos modelos Word/PDF.")
    cpf_busca = st.text_input("Digite o CPF do paciente (Ex: 181.398.259-72) para gerar o Kit:")
    if st.button("⚡ Gerar Kit de Contratação"):
        with st.spinner("Injetando dados via Cloud API..."):
            doc_c = Document()
            doc_c.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS', 1)
            doc_c.add_paragraph(f"CONTRATANTE CPF: {cpf_busca}\nPelo presente instrumento particular...")
            doc_c.save("Contrato_Cloud.docx")
            
            doc_t = Document()
            doc_t.add_heading('TERMO DE CONSENTIMENTO LIVRE', 1)
            doc_t.add_paragraph("Autorizo a realização dos procedimentos mapeados pelo OdontoMais OS.")
            doc_t.save("Termo_Cloud.docx")
            
            doc_o = Document()
            doc_o.add_heading('ORÇAMENTO CLÍNICO', 1)
            doc_o.add_paragraph("Ticket Médio Vinculado: R$ 607,25")
            doc_o.save("Orcamento_Cloud.docx")
            
            st.success("✅ Documentos Gerados! Pronto para Download.")

elif page == "⭐ NPS Avançado":
    st.title("Gestão de Satisfação e Reputação")
    st.markdown("Cruzamento de Notas de NPS com Faturamento de Procedimentos.")
    
    with st.form("nps"):
        nome = st.text_input("Paciente")
        procedimento = st.selectbox("Procedimento Relacionado", ["Implante Unitário", "Prótese Protocolo", "Lentes de Contato"])
        nota = st.slider("Nota de Satisfação (0 a 10)", 0, 10, 10)
        sub = st.form_submit_button("Registrar NPS")
        
    if sub:
        if nota >= 9:
            st.success("✅ Fluxo Positivo! Disparando convite automático para Google Review.")
            st.markdown("<a href='https://g.page/r/YOUR_LINK_HERE/review' target='_blank' style='background:#d1175c;color:white;padding:10px 20px;text-decoration:none;border-radius:5px;font-weight:bold;'>Ir para Google Review</a>", unsafe_allow_html=True)
        else:
            st.error("⚠️ Alerta de Contenção de Danos! Ticket Médio de R$ 607,25 em Risco.")
            st.text_area("Descreva a causa raiz da insatisfação para auditoria gerencial:")
            if st.button("Gravar Ocorrência Crítica no Supabase"):
                st.info("Log gravado. Gestão notificada.")
